from docx import Document as DocxDocument

def save_rag_output_to_docx(output_filename: str, user_question: str, llm_final_answer: str, context_docs_used: list):
    """
    Saves the RAG pipeline output (query, answer, and sources) to a DOCx file.
    """
    if not (llm_final_answer and isinstance(llm_final_answer, str) and llm_final_answer not in ["Error generating response from LLM.", "LLM is not initialized.", "RAG chain skipped: Dependencies not met.", "RAG chain skipped: Collection not found."]):
        print("Skipping DOCx saving: No valid LLM response to save.")
        return

    try:
        document = DocxDocument()
        document.add_heading('RAG Pipeline Output', level=1)

        document.add_paragraph("Query:")
        document.add_paragraph(user_question)

        document.add_heading('Answer', level=2)
        document.add_paragraph(llm_final_answer)

        if context_docs_used:
            document.add_heading('Sources Used for Context', level=2)
            for i, doc in enumerate(context_docs_used):
                source_info_parts = []
                source_info_parts.append(f"Document {i+1}:")
                source_info_parts.append(f"Source: {doc.metadata.get('source', 'N/A')}")
                source_info_parts.append(f"Page: {doc.metadata.get('page', 'N/A')}")
                if doc.metadata.get('type') == 'table' and doc.metadata.get('table_num') is not None:
                     source_info_parts.append(f"Table: {doc.metadata.get('table_num')}")

                source_info = ", ".join(source_info_parts)
                document.add_paragraph(source_info)

        document.save(output_filename)
        print(f"LLM output saved to '{output_filename}'.")
    except Exception as e:
        print(f"Error saving to DOCx: {e}")
        print("Please ensure 'python-docx' is installed.")